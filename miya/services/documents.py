"""Local document text extraction (spec §6).

Documents shared in monitored chats — invoices, packing lists, price sheets —
are parsed **on the VPS**, never uploaded anywhere. Only the extracted text
goes to the extractor, and only up to ``DOC_MAX_CHARS`` of it.

Every parser is best-effort: an encrypted PDF, a corrupt workbook or an
unsupported legacy format returns ``None`` and the caller falls back to
metadata + caption, exactly like a video.
"""

from __future__ import annotations

import asyncio
import csv
import io
import logging
from dataclasses import dataclass
from pathlib import Path

from miya.config import settings

log = logging.getLogger(__name__)

# What we can read locally. Legacy binary .xls needs xlrd and is deliberately
# not supported — it falls through to metadata-only rather than pulling in
# another dependency for a format the owner rarely receives.
SUPPORTED_SUFFIXES = {".pdf", ".xlsx", ".xlsm", ".docx", ".txt", ".csv", ".md"}

TRUNCATION_NOTE = "\n\n[… hujjat qisqartirildi]"


@dataclass(slots=True)
class DocumentText:
    text: str
    truncated: bool = False
    detail: str | None = None  # page/sheet count, for the interaction metadata


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_SUFFIXES


def _cap(text: str) -> DocumentText:
    text = text.strip()
    if len(text) <= settings.doc_max_chars:
        return DocumentText(text=text)
    return DocumentText(
        text=text[: settings.doc_max_chars] + TRUNCATION_NOTE, truncated=True
    )


def _read_pdf(path: Path) -> DocumentText:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
            # Stop early on a 300-page catalogue: the cap would drop the rest
            # anyway and every extra page costs seconds of CPU.
            if sum(len(p) for p in pages) > settings.doc_max_chars:
                break
        total = len(pdf.pages)
    result = _cap("\n\n".join(pages))
    result.detail = f"{total} sahifa"
    return result


def _read_xlsx(path: Path) -> DocumentText:
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    try:
        chunks: list[str] = []
        for sheet in workbook.worksheets:
            chunks.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    chunks.append(" | ".join(cells))
                if sum(len(c) for c in chunks) > settings.doc_max_chars:
                    break
        sheets = len(workbook.worksheets)
    finally:
        workbook.close()
    result = _cap("\n".join(chunks))
    result.detail = f"{sheets} varaq"
    return result


def _read_docx(path: Path) -> DocumentText:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _cap("\n".join(parts))


def _read_text(path: Path) -> DocumentText:
    raw = path.read_bytes()[: settings.doc_max_chars * 4]
    text = raw.decode("utf-8", errors="replace")
    if path.suffix.lower() == ".csv":
        rows = list(csv.reader(io.StringIO(text)))
        text = "\n".join(" | ".join(cell for cell in row if cell) for row in rows)
    return _cap(text)


_READERS = {
    ".pdf": _read_pdf,
    ".xlsx": _read_xlsx,
    ".xlsm": _read_xlsx,
    ".docx": _read_docx,
    ".txt": _read_text,
    ".csv": _read_text,
    ".md": _read_text,
}


def read_document(path: Path) -> DocumentText | None:
    """Parse one document. Returns None when it cannot be read locally."""
    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        return None
    try:
        result = reader(path)
    except Exception:
        # Encrypted PDFs, corrupt workbooks, mislabelled extensions — all end
        # up here and are treated exactly like an unsupported format.
        log.warning("could not read document %s", path.name, exc_info=True)
        return None
    return result if result.text else None


async def read_document_async(path: Path) -> DocumentText | None:
    """Off the event loop: PDF parsing is CPU-bound and can take seconds."""
    return await asyncio.to_thread(read_document, path)
