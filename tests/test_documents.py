"""Local document parsing (spec §6). Real files, real parsers, no network."""

from __future__ import annotations

from miya.config import settings
from miya.services import documents


def test_a_plain_text_file_is_read_as_is(tmp_path):
    path = tmp_path / "eslatma.txt"
    path.write_text("Akmalga 5 mln berdim", encoding="utf-8")

    parsed = documents.read_document(path)

    assert parsed.text == "Akmalga 5 mln berdim"
    assert parsed.truncated is False


def test_a_csv_becomes_readable_rows(tmp_path):
    path = tmp_path / "narxlar.csv"
    path.write_text("mahsulot,narx\nchoy,25000\nqand,18000\n", encoding="utf-8")

    parsed = documents.read_document(path)

    assert "mahsulot | narx" in parsed.text
    assert "choy | 25000" in parsed.text


def test_an_xlsx_is_flattened_sheet_by_sheet(tmp_path):
    from openpyxl import Workbook

    path = tmp_path / "hisob.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Avgust"
    sheet.append(["mijoz", "summa"])
    sheet.append(["Akmal", 5000000])
    workbook.save(path)

    parsed = documents.read_document(path)

    assert "# Avgust" in parsed.text
    assert "Akmal | 5000000" in parsed.text
    assert parsed.detail == "1 varaq"


def test_a_docx_keeps_paragraphs_and_table_cells(tmp_path):
    import docx

    path = tmp_path / "shartnoma.docx"
    document = docx.Document()
    document.add_paragraph("Yuk tashish shartnomasi")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Summa"
    table.rows[0].cells[1].text = "12000 USD"
    document.save(path)

    parsed = documents.read_document(path)

    assert "Yuk tashish shartnomasi" in parsed.text
    assert "Summa | 12000 USD" in parsed.text


def test_long_documents_are_capped_and_marked(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "doc_max_chars", 100)
    path = tmp_path / "katalog.txt"
    path.write_text("x" * 5000, encoding="utf-8")

    parsed = documents.read_document(path)

    assert parsed.truncated is True
    assert parsed.text.endswith(documents.TRUNCATION_NOTE)
    assert len(parsed.text) < 200


def test_a_corrupt_file_returns_nothing_instead_of_raising(tmp_path):
    path = tmp_path / "buzilgan.pdf"
    path.write_bytes(b"this is definitely not a pdf")

    assert documents.read_document(path) is None


def test_an_unsupported_extension_is_declined_up_front(tmp_path):
    path = tmp_path / "arxiv.zip"
    path.write_bytes(b"PK\x03\x04")

    assert documents.is_supported(path) is False
    assert documents.read_document(path) is None


def test_an_empty_document_counts_as_unreadable(tmp_path):
    path = tmp_path / "bosh.txt"
    path.write_text("   \n\n", encoding="utf-8")

    assert documents.read_document(path) is None


async def test_the_async_wrapper_returns_the_same_result(tmp_path):
    path = tmp_path / "eslatma.txt"
    path.write_text("bojxona to'lovi 3 mln", encoding="utf-8")

    parsed = await documents.read_document_async(path)

    assert parsed.text == "bojxona to'lovi 3 mln"
